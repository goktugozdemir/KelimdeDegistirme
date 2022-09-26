import os
from tkinter import *
from tkinter import ttk
from docx import Document


taraf = []


def rClicker(e):
    try:
        def rClick_Copy(e, apnd=0):
            e.widget.event_generate('<Control-c>')

        def rClick_Cut(e):
            e.widget.event_generate('<Control-x>')

        def rClick_Paste(e):
            e.widget.event_generate('<Control-v>')

        e.widget.focus()

        nclst = [
            (' Kes', lambda e=e: rClick_Cut(e)),
            (' Kopyala', lambda e=e: rClick_Copy(e)),
            (' Yapıştır', lambda e=e: rClick_Paste(e)),
        ]

        rmenu = Menu(None, tearoff=0, takefocus=0)

        for (txt, cmd) in nclst:
            rmenu.add_command(label=txt, command=cmd)

        rmenu.tk_popup(e.x_root + 40, e.y_root + 10, entry="0")

    except TclError:
        print(' - rClick menu, something wrong')
        pass

    return "break"


def rClickbinder(r):
    try:
        for b in ['Text', 'Entry', 'Listbox', 'Label']:  #
            r.bind_class(b, sequence='<Button-3>',
                         func=rClicker, add='')
    except TclError:
        print(' - rClickbinder, something wrong')
        pass


def display_text():
    global entry
    global string1
    global string2
    global string3
    global string4
    global string5
    string1 = yazı.get()
    string2 = yazı2.get()
    string3 = yazı3.get()
    string4 = yazı4.get()
    string5 = yazı5.get()
    window.destroy()


def display_text1():
    global tarafs
    taraf.append(t1.get())
    taraf.append(t2.get())
    taraf.append(t3.get())
    taraf.append(t4.get())
    taraf.append(t5.get())
    taraf.append(t6.get())
    taraf.append(t7.get())
    taraf.append(t8.get())
    tarafs = tarafs + 1
    t1.delete(0, END)
    t1.insert(0, "")
    t2.delete(0, END)
    t2.insert(0, "")
    t3.delete(0, END)
    t3.insert(0, "")
    t4.delete(0, END)
    t4.insert(0, "")
    t5.delete(0, END)
    t5.insert(0, "")
    t6.delete(0, END)
    t6.insert(0, "")
    t7.delete(0, END)
    t7.insert(0, "")
    t8.delete(0, END)
    t8.insert(0, "")


tarafs = 0


def display_text2():
    global aaa1
    global aaa2
    global aaa3
    global aaa4
    global aaa5
    aaa1 = gg.get()
    aaa2 = gg2.get()
    aaa3 = gg3.get()
    aaa4 = gg4.get()
    aaa5 = gg5.get()
    window2.destroy()


def display_text3():
    global dosyasss
    dosyasss = dosya1.get()
    window3.destroy()


def dur():
    window1.destroy()



def quit_loop():
    var.get()
    global selection
    selection = var.get()
    window0.destroy()


window0 = Tk()
width = window0.winfo_screenwidth()
height = window0.winfo_screenheight()
global screen_resolution
screen_resolution = str(width) + 'x' + str(height)
var = IntVar()
var.set(1)
window0.geometry(screen_resolution)
window0.option_add("*font", "lucida 25 bold italic")
window0.configure(background="white");
sssss = Label(window0, text="Uyuşmazlık Türünü Seçiniz", justify=LEFT, anchor="w")
sssss.pack(side='top')
R1 = Radiobutton(window0, text="Ticari", variable=var, value=1)
R1.pack(anchor = W )
R2 = Radiobutton(window0, text="İş Hukuku", variable=var, value=2)
R2.pack(anchor = W )
R3 = Radiobutton(window0, text="Tüketici", variable=var, value=3)
R3.pack(anchor = W)
btn = Button(window0, text = 'Onayla', bd = '3',command = quit_loop)
btn.pack(side='top')
window0.mainloop()

current=os.getcwd()
if(selection==1):
    future=current+r"\ticari"
    os.chdir(future)
elif(selection==2):
    future = current + r"\iş Hukuku"
    os.chdir(future)
else:
    future = current + r"\tüketici"
    os.chdir(future)

window = Tk()
window.geometry(screen_resolution)
window.option_add("*font", "lucida 25 bold italic")
window.configure(background="white");
a = Label(window, text="Arabuluculuk Bürosu", justify=LEFT, anchor="w").grid(sticky="W", row=0, column=0)
b = Label(window, text="Büro Dosya Numarası", justify=LEFT, anchor="w").grid(sticky="W", row=1, column=0)
c = Label(window, text="Arabuluculuk Numarası", justify=LEFT, anchor="w").grid(sticky="W", row=2, column=0)
d = Label(window, text="Arabuluculuk Başvuru Tarihi", justify=LEFT, anchor="w").grid(sticky="W", row=3, column=0)
e = Label(window, text="Arabuluculuk Konusu Uyuşmazlık ", justify=LEFT, anchor="w").grid(sticky="W", row=4, column=0)
yazı = Entry(window, width=40)
yazı.bind('<Button-3>', rClicker, add='')
yazı.grid(row=0, column=1)
yazı.focus_set()
yazı2 = Entry(window, width=40)
yazı2.grid(row=1, column=1)
yazı2.bind('<Button-3>', rClicker, add='')
yazı2.focus_set()
yazı3 = Entry(window, width=40)
yazı3.grid(row=2, column=1)
yazı3.bind('<Button-3>', rClicker, add='')
yazı3.focus_set()
yazı4 = Entry(window, width=40)
yazı4.grid(row=3, column=1)
yazı4.bind('<Button-3>', rClicker, add='')
yazı4.focus_set()
yazı5 = Entry(window, width=40)
yazı5.bind('<Button-3>', rClicker, add='')
yazı5.grid(row=4, column=1)
yazı5.focus_set()
Button = ttk.Button(window, text="Taraf Bilgisi Adımına Geçiniz", width=50, command=display_text)
Button.grid(row=5, column=1)
window.mainloop()

window1 = Tk()
window1.geometry(screen_resolution)
window1.option_add("*font", "lucida 25 bold italic")
window1.configure(background="white");
a1 = Label(window1, text="Ad ve Soyad/Ünvan", justify=LEFT, anchor="w").grid(sticky="W", row=0, column=0)
b1 = Label(window1, text="Ticaret Sicil Numarası/T.C Kimlik Numarası", justify=LEFT, anchor="w").grid(sticky="W", row=1,
                                                                                                      column=0)
c1 = Label(window1, text="Adresi", anchor="w").grid(sticky="W", row=2, column=0)
c1 = Label(window1, text="Asil Telefon", justify=LEFT, anchor="w").grid(sticky="W", row=3, column=0)
d1 = Label(window1, text="Vekil", justify=LEFT, anchor="w").grid(sticky="W", row=4, column=0)
e1 = Label(window1, text="T.C. Kimlik Numarası", justify=LEFT, anchor="w").grid(sticky="W", row=5, column=0)
f1 = Label(window1, text="Vekil Telefon Numarası", justify=LEFT, anchor="w").grid(sticky="W", row=6, column=0)
g1 = Label(window1, text="E-mail", justify=LEFT, anchor="w").grid(sticky="W", row=7, column=0)
t1 = Entry(window1, width=40)
t1.bind('<Button-3>', rClicker, add='')
t1.grid(row=0, column=1)
t1.focus_set()
t2 = Entry(window1, width=40)
t2.bind('<Button-3>', rClicker, add='')
t2.grid(row=1, column=1)
t2.focus_set()
t3 = Entry(window1, width=40)
t3.bind('<Button-3>', rClicker, add='')
t3.grid(row=2, column=1)
t3.focus_set()
t4 = Entry(window1, width=40)
t4.bind('<Button-3>', rClicker, add='')
t4.grid(row=3, column=1)
t4.focus_set()
t5 = Entry(window1, width=40)
t5.bind('<Button-3>', rClicker, add='')
t5.grid(row=4, column=1)
t5.focus_set()
t6 = Entry(window1, width=40)
t6.bind('<Button-3>', rClicker, add='')
t6.grid(row=5, column=1)
t6.focus_set()
t7 = Entry(window1, width=40)
t7.bind('<Button-3>', rClicker, add='')
t7.grid(row=6, column=1)
t7.focus_set()
t8 = Entry(window1, width=40)
t8.bind('<Button-3>', rClicker, add='')
t8.grid(row=7, column=1)
t8.focus_set()
Button1 = ttk.Button(window1, text="Sıradaki Tarafı Giriniz", width=25, command=display_text1)
Button1.grid(row=8, column=1)
Button2 = ttk.Button(window1, text="Taraf Bilgilerini Girmeyi Bitir", width=25, command=dur)
Button2.grid(row=9, column=1)
tttt = Label(window1, text="Boş Bırakılacak yerlere nokta koyunuz.").grid(sticky="W", row=8, column=0)
window.mainloop()

window2 = Tk()
window2.geometry(screen_resolution)
window2.option_add("*font", "lucida 25 bold italic")
window2.configure(background="white");
q = Label(window2, text="Arabuluculuk Bürosuna Başvuru Tarihi ", justify=LEFT, anchor="w").grid(sticky="W", row=0,
                                                                                                column=0)
w = Label(window2, text="Arabulucunun Görevlendirildiği Tarih", justify=LEFT, anchor="w").grid(sticky="W", row=1,
                                                                                               column=0)
j = Label(window2, text="Arabuluculuk Sürecinin Bittiği Tarih ", justify=LEFT, anchor="w").grid(sticky="W", row=2,
                                                                                                column=0)
k = Label(window2, text="Son Tutanağın Düzenlendiği Yer", justify=LEFT, anchor="w").grid(sticky="W", row=3, column=0)
l = Label(window2, text="Son Tutanağın Düzenlendiği Tarih", justify=LEFT, anchor="w").grid(sticky="W", row=4, column=0)
gg = Entry(window2, width=40)
gg.bind('<Button-3>', rClicker, add='')
gg.grid(row=0, column=1)
gg.focus_set()
gg2 = Entry(window2, width=40)
gg2.bind('<Button-3>', rClicker, add='')
gg2.grid(row=1, column=1)
gg2.focus_set()
gg3 = Entry(window2, width=40)
gg3.bind('<Button-3>', rClicker, add='')
gg3.grid(row=2, column=1)
gg3.focus_set()
gg4 = Entry(window2, width=40)
gg4.bind('<Button-3>', rClicker, add='')
gg4.grid(row=3, column=1)
gg4.focus_set()
gg5 = Entry(window2, width=40)
gg5.bind('<Button-3>', rClicker, add='')
gg5.grid(row=4, column=1)
gg5.focus_set()
Button = ttk.Button(window2, text="Bilgileri Girdiyseniz Basın", width=50, command=display_text2)
Button.grid(row=5, column=1)
tarafs = Label(window2, text=tarafs).grid(row=6, column=1)
window2.mainloop()

window3 = Tk()
window3.geometry(screen_resolution)
window3.configure(background="white");
dosya = Label(window3, text="Dosyalar Masaüstündeki Hangi klasöre kaydedilsin", justify=LEFT, anchor="w").grid(
    sticky="W", row=0, column=0)
dosya1 = Entry(window3, width=40)
dosya1.grid(row=0, column=1)
dosya1.focus_set()
Button = ttk.Button(window3, text="Onayla", width=50, command=display_text3)
Button.grid(row=5, column=1)
window3.mainloop()
case=0
tarafsayı=len(taraf)
print(tarafsayı)
desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
dosyayolu = ((desktop +'//Arabuluculuk Evrak'+ dosyasss))
list = ['B1', 'B2', 'B3', 'B4', 'B5', 'G1', 'G2', 'G3', 'G4', 'G5', 'T11', 'T12', 'T13', 'T14', 'T15', 'T16', 'T17',
        'T18', 'T21', 'T22', 'T23', 'T24', 'T25', 'T26', 'T27', 'T28']
list2 = [string1, string2, string3, string4, string5, aaa1, aaa2, aaa3, aaa4, aaa5, taraf[0], taraf[1], taraf[2],
         taraf[3], taraf[4], taraf[5], taraf[6], taraf[7], taraf[8], taraf[9], taraf[10], taraf[11], taraf[12],
         taraf[13], taraf[14], taraf[15]]
if (tarafsayı > 73 and  case==0):
    list = ['B1', 'B2', 'B3', 'B4', 'B5', 'G1', 'G2', 'G3', 'G4', 'G5', 'T11', 'T12', 'T13', 'T14', 'T15', 'T16', 'T17',
            'T18', 'T21', 'T22', 'T23', 'T24', 'T25', 'T26', 'T27', 'T28', 'T31', 'T32', 'T33', 'T34', 'T35', 'T36',
            'T37', 'T38', 'T41', 'T42', 'T43', 'T44', 'T45', 'T46', 'T47', 'T48', 'T51', 'T52', 'T53', 'T54', 'T55'
        , 'T56', 'T57', 'T58', 'T61', 'T62', 'T63', 'T64', 'T65', 'T66', 'T67', 'T68', 'T71', 'T72', 'T73', 'T74',
            'T75', 'T76', 'T77', 'T78'
        , 'T81', 'T82', 'T83', 'T84', 'T85', 'T86', 'T87', 'T91', 'T92', 'T93', 'T94', 'T95', 'T96', 'T97', 'T98', 'T101'
            , 'T102', 'T103', 'T104', 'T105', 'T106', 'T107', 'T108']

    list2 = [string1, string2, string3, string4, string5, aaa1, aaa2, aaa3, aaa4, aaa5, taraf[0], taraf[1], taraf[2],
             taraf[3], taraf[4], taraf[5], taraf[6], taraf[7], taraf[8], taraf[9], taraf[10], taraf[11], taraf[12],
             taraf[13], taraf[14], taraf[15], taraf[16], taraf[17], taraf[18], taraf[19], taraf[20], taraf[21],
             taraf[22], taraf[23], taraf[24]
        , taraf[25], taraf[26], taraf[27], taraf[28], taraf[29], taraf[30], taraf[31], taraf[32], taraf[33], taraf[34],
             taraf[35]
        , taraf[36], taraf[37], taraf[38], taraf[39], taraf[40], taraf[41], taraf[42], taraf[43]
        , taraf[44], taraf[45], taraf[46], taraf[47], taraf[48], taraf[49], taraf[50], taraf[51], taraf[52], taraf[53],
             taraf[54], taraf[55]
        , taraf[56], taraf[57], taraf[58], taraf[59], taraf[60], taraf[61], taraf[62], taraf[63], taraf[64], taraf[65],
             taraf[66], taraf[67], taraf[68]
        , taraf[69], taraf[70], taraf[71], taraf[72], taraf[73], taraf[74], taraf[75], taraf[76], taraf[77], taraf[78], taraf[79]]
    case = 1
if (tarafsayı > 65 and  case == 0):
        list = ['B1', 'B2', 'B3', 'B4', 'B5', 'G1', 'G2', 'G3', 'G4', 'G5', 'T11', 'T12', 'T13', 'T14', 'T15', 'T16',
                'T17',
                'T18', 'T21', 'T22', 'T23', 'T24', 'T25', 'T26', 'T27', 'T28', 'T31', 'T32', 'T33', 'T34', 'T35', 'T36',
                'T37', 'T38', 'T41', 'T42', 'T43', 'T44', 'T45', 'T46', 'T47', 'T48', 'T51', 'T52', 'T53', 'T54', 'T55'
            , 'T56', 'T57', 'T58', 'T61', 'T62', 'T63', 'T64', 'T65', 'T66', 'T67', 'T68', 'T71', 'T72', 'T73', 'T74',
                'T75', 'T76', 'T77', 'T78'
            , 'T81', 'T82', 'T83', 'T84', 'T85', 'T86', 'T87', 'T91', 'T92', 'T93', 'T94', 'T95', 'T96', 'T97', 'T98']

        list2 = [string1, string2, string3, string4, string5, aaa1, aaa2, aaa3, aaa4, aaa5, taraf[0], taraf[1],
                 taraf[2],
                 taraf[3], taraf[4], taraf[5], taraf[6], taraf[7], taraf[8], taraf[9], taraf[10], taraf[11], taraf[12],
                 taraf[13], taraf[14], taraf[15], taraf[16], taraf[17], taraf[18], taraf[19], taraf[20], taraf[21],
                 taraf[22], taraf[23], taraf[24]
            , taraf[25], taraf[26], taraf[27], taraf[28], taraf[29], taraf[30], taraf[31], taraf[32], taraf[33],
                 taraf[34],
                 taraf[35]
            , taraf[36], taraf[37], taraf[38], taraf[39], taraf[40], taraf[41], taraf[42], taraf[43]
            , taraf[44], taraf[45], taraf[46], taraf[47], taraf[48], taraf[49], taraf[50], taraf[51], taraf[52],
                 taraf[53],
                 taraf[54], taraf[55]
            , taraf[56], taraf[57], taraf[58], taraf[59], taraf[60], taraf[61], taraf[62], taraf[63], taraf[64],
                 taraf[65], taraf[66], taraf[67], taraf[68]
            , taraf[69], taraf[70], taraf[71]]
        case = 2

if (tarafsayı > 57 and  case == 0):
    list = ['B1', 'B2', 'B3', 'B4', 'B5', 'G1', 'G2', 'G3', 'G4', 'G5', 'T11', 'T12', 'T13', 'T14', 'T15', 'T16', 'T17',
            'T18', 'T21', 'T22', 'T23', 'T24', 'T25', 'T26', 'T27', 'T28', 'T31', 'T32', 'T33', 'T34', 'T35', 'T36',
            'T37', 'T38', 'T41', 'T42', 'T43', 'T44', 'T45', 'T46', 'T47', 'T48', 'T51', 'T52', 'T53', 'T54', 'T55'
        , 'T56', 'T57', 'T58', 'T61', 'T62', 'T63', 'T64', 'T65', 'T66', 'T67', 'T68', 'T71', 'T72', 'T73', 'T74',
            'T75', 'T76', 'T77', 'T78'
        , 'T81', 'T82', 'T83', 'T84', 'T85', 'T86', 'T87', 'T88']

    list2 = [string1, string2, string3, string4, string5, aaa1, aaa2, aaa3, aaa4, aaa5, taraf[0], taraf[1], taraf[2],
             taraf[3], taraf[4], taraf[5], taraf[6], taraf[7], taraf[8], taraf[9], taraf[10], taraf[11], taraf[12],
             taraf[13], taraf[14], taraf[15], taraf[16], taraf[17], taraf[18], taraf[19], taraf[20], taraf[21],
             taraf[22], taraf[23], taraf[24]
        , taraf[25], taraf[26], taraf[27], taraf[28], taraf[29], taraf[30], taraf[31], taraf[32], taraf[33], taraf[34],
             taraf[35]
        , taraf[36], taraf[37], taraf[38], taraf[39], taraf[40], taraf[41], taraf[42], taraf[43]
        , taraf[44], taraf[45], taraf[46], taraf[47], taraf[48], taraf[49], taraf[50], taraf[51], taraf[52], taraf[53],
             taraf[54], taraf[55]
        , taraf[56], taraf[57], taraf[58], taraf[59], taraf[60], taraf[61], taraf[62], taraf[63]]
    case = 3
if (tarafsayı > 49 and  case==0):
    list = ['B1', 'B2', 'B3', 'B4', 'B5', 'G1', 'G2', 'G3', 'G4', 'G5', 'T11', 'T12', 'T13', 'T14', 'T15', 'T16', 'T17',
            'T18', 'T21', 'T22', 'T23', 'T24', 'T25', 'T26', 'T27', 'T28', 'T31', 'T32', 'T33', 'T34', 'T35', 'T36',
            'T37', 'T38', 'T41', 'T42', 'T43', 'T44', 'T45', 'T46', 'T47', 'T48', 'T51', 'T52', 'T53', 'T54', 'T55'
        , 'T56', 'T57', 'T58', 'T61', 'T62', 'T63', 'T64', 'T65', 'T66', 'T67', 'T68', 'T71', 'T72', 'T73', 'T74',
            'T75', 'T76', 'T77', 'T78']
    list2 = [string1, string2, string3, string4, string5, aaa1, aaa2, aaa3, aaa4, aaa5, taraf[0], taraf[1], taraf[2],
             taraf[3], taraf[4], taraf[5], taraf[6], taraf[7], taraf[8], taraf[9], taraf[10], taraf[11], taraf[12],
             taraf[13], taraf[14], taraf[15], taraf[16], taraf[17], taraf[18], taraf[19], taraf[20], taraf[21],
             taraf[22], taraf[23], taraf[24]
        , taraf[25], taraf[26], taraf[27], taraf[28], taraf[29], taraf[30], taraf[31], taraf[32], taraf[33], taraf[34],
             taraf[35]
        , taraf[36], taraf[37], taraf[38], taraf[39], taraf[40], taraf[41], taraf[42], taraf[43]
        , taraf[44], taraf[45], taraf[46], taraf[47], taraf[48], taraf[49], taraf[50], taraf[51], taraf[52], taraf[53],
             taraf[54], taraf[55]]
    case = 4
if (tarafsayı > 41 and  case==0):
    list = ['B1', 'B2', 'B3', 'B4', 'B5', 'G1', 'G2', 'G3', 'G4', 'G5', 'T11', 'T12', 'T13', 'T14', 'T15', 'T16', 'T17',
            'T18', 'T21', 'T22', 'T23', 'T24', 'T25', 'T26', 'T27', 'T28', 'T31', 'T32', 'T33', 'T34', 'T35', 'T36',
            'T37', 'T38', 'T41', 'T42', 'T43', 'T44', 'T45', 'T46', 'T47', 'T48', 'T51', 'T52', 'T53', 'T54', 'T55'
        , 'T56', 'T57', 'T58', 'T61', 'T62', 'T63', 'T64', 'T65', 'T66', 'T67', 'T68']

    list2 = [string1, string2, string3, string4, string5, aaa1, aaa2, aaa3, aaa4, aaa5, taraf[0], taraf[1], taraf[2],
             taraf[3], taraf[4], taraf[5], taraf[6], taraf[7], taraf[8], taraf[9], taraf[10], taraf[11], taraf[12],
             taraf[13], taraf[14], taraf[15], taraf[16], taraf[17], taraf[18], taraf[19], taraf[20], taraf[21],
             taraf[22], taraf[23], taraf[24]
        , taraf[25], taraf[26], taraf[27], taraf[28], taraf[29], taraf[30], taraf[31], taraf[32], taraf[33], taraf[34],
             taraf[35]
        , taraf[36], taraf[37], taraf[38], taraf[39], taraf[40], taraf[41], taraf[42], taraf[43]
        , taraf[44], taraf[45], taraf[46], taraf[47]]
    case = 5
if (tarafsayı > 33 and  case==0):
    list = ['B1', 'B2', 'B3', 'B4', 'B5', 'G1', 'G2', 'G3', 'G4', 'G5', 'T11', 'T12', 'T13', 'T14', 'T15', 'T16', 'T17',
            'T18', 'T21', 'T22', 'T23', 'T24', 'T25', 'T26', 'T27', 'T28', 'T31', 'T32', 'T33', 'T34', 'T35', 'T36',
            'T37', 'T38', 'T41', 'T42', 'T43', 'T44', 'T45', 'T46', 'T47', 'T48', 'T51', 'T52', 'T53', 'T54', 'T55'
        , 'T56', 'T57', 'T58']
    list2 = [string1, string2, string3, string4, string5, aaa1, aaa2, aaa3, aaa4, aaa5, taraf[0], taraf[1], taraf[2],
             taraf[3], taraf[4], taraf[5], taraf[6], taraf[7], taraf[8], taraf[9], taraf[10], taraf[11], taraf[12],
             taraf[13], taraf[14], taraf[15], taraf[16], taraf[17], taraf[18], taraf[19], taraf[20], taraf[21],
             taraf[22], taraf[23], taraf[24]
        , taraf[25], taraf[26], taraf[27], taraf[28], taraf[29], taraf[30], taraf[31], taraf[32], taraf[33], taraf[34],
             taraf[35]
        , taraf[36], taraf[37], taraf[38], taraf[39]]
    case = 6
if(tarafsayı>25 and  case==0):
    list = ['B1', 'B2', 'B3', 'B4', 'B5', 'G1', 'G2', 'G3', 'G4', 'G5', 'T11', 'T12', 'T13', 'T14', 'T15', 'T16', 'T17',
            'T18', 'T21', 'T22', 'T23', 'T24', 'T25', 'T26', 'T27', 'T28', 'T31', 'T32', 'T33', 'T34', 'T35', 'T36',
            'T37', 'T38', 'T41', 'T42', 'T43', 'T44', 'T45', 'T46', 'T47', 'T48']
    list2 =[ string1, string2, string3, string4, string5, aaa1, aaa2, aaa3, aaa4, aaa5, taraf[0], taraf[1], taraf[2],
    taraf[3], taraf[4], taraf[5], taraf[6], taraf[7], taraf[8], taraf[9], taraf[10], taraf[11], taraf[12],
    taraf[13], taraf[14], taraf[15], taraf[16], taraf[17], taraf[18], taraf[19], taraf[20], taraf[21],
    taraf[22], taraf[23], taraf[24]
, taraf[25], taraf[26], taraf[27], taraf[28], taraf[29], taraf[30], taraf[31]]
    case = 7
if(tarafsayı>17 and  case==0):
    list = ['B1', 'B2', 'B3', 'B4', 'B5', 'G1', 'G2', 'G3', 'G4', 'G5', 'T11', 'T12', 'T13', 'T14', 'T15', 'T16', 'T17',
            'T18', 'T21', 'T22', 'T23', 'T24', 'T25', 'T26', 'T27', 'T28', 'T31', 'T32', 'T33', 'T34', 'T35', 'T36',
            'T37', 'T38']
    list2 = [string1, string2, string3, string4, string5, aaa1, aaa2, aaa3, aaa4, aaa5, taraf[0], taraf[1], taraf[2],
         taraf[3], taraf[4], taraf[5], taraf[6], taraf[7], taraf[8], taraf[9], taraf[10], taraf[11], taraf[12],
         taraf[13], taraf[14], taraf[15], taraf[16], taraf[17], taraf[18], taraf[19], taraf[20], taraf[21], taraf[22], taraf[23]]
    case=8



liste = ['B1', 'B2', 'B3', 'B4', 'B5', 'G1', 'G2', 'G3', 'G4', 'G5', 'T11', 'T12', 'T13', 'T14', 'T15', 'T16', 'T17',
            'T18', 'T21', 'T22', 'T23', 'T24', 'T25', 'T26', 'T27', 'T28', 'T31', 'T32', 'T33', 'T34', 'T35', 'T36',
            'T37', 'T38', 'T41', 'T42', 'T43', 'T44', 'T45', 'T46', 'T47', 'T48', 'T51', 'T52', 'T53', 'T54', 'T55'
        , 'T56', 'T57', 'T58', 'T61', 'T62', 'T63', 'T64', 'T65', 'T66', 'T67', 'T68', 'T71', 'T72', 'T73', 'T74',
            'T75', 'T76', 'T77', 'T78'
        , 'T81', 'T82', 'T83', 'T84', 'T85', 'T86', 'T87', 'T91', 'T92', 'T93', 'T94', 'T95', 'T96', 'T97', 'T98', 'T101'
            , 'T102', 'T103', 'T104', 'T105', 'T106', 'T107', 'T108']
if(case==0):
 for i in range(3,9,1):
    a=str(i)
    liste.append("Taraf"+" "+a)
else:
    for i in range(11-case,9, +1):
        a = str(i)
        liste.append("Taraf" + " " + a)

print(liste)
list(set(liste) - set(list))

def replace_string1(filename='1.DAVET MEKTUBU.docx'):
    doc = Document(filename)
    for p in doc.paragraphs:
        inline = p.runs
        for j in range(0, len(inline)):
            for i in range(0, len(list)):
               p.text = p.text.replace(list[i], list2[i])
    doc.save(dosyayolu + '1.doc')
    return 1

def replace_string2(filename='2.BİLGİLENDİRME TUTANAĞI.docx'):
    doc = Document(filename)
    for p in doc.paragraphs:
        inline = p.runs
        for j in range(0, len(inline)):
            for i in range(0, len(list)):
               p.text = p.text.replace(list[i], list2[i])
    doc.save(dosyayolu + '2.doc')
    return 1


def replace_string3(filename='3.İLK OTURUM TUTANAĞI.docx'):
    doc = Document(filename)
    for p in doc.paragraphs:
        inline = p.runs
        for j in range(0, len(inline)):
            for i in range(0, len(list)):
               p.text = p.text.replace(list[i], list2[i])
    doc.save(dosyayolu + '3.doc')
    return 1


def replace_string4(filename='4.İKİNCİ OTURUM TUTANAĞI.docx'):
 doc = Document(filename)
 for p in doc.paragraphs:
    inline = p.runs
    for j in range(0, len(inline)):
        for i in range(0, len(list)):
           p.text = p.text.replace(list[i], list2[i])
 doc.save(dosyayolu + '4.doc')
 return 1


def replace_string5(filename='5.ANLAŞMA BELGESİ.docx'):
 doc = Document(filename)
 for p in doc.paragraphs:
    inline = p.runs
    for j in range(0, len(inline)):
        for i in range(0, len(list)):
           p.text = p.text.replace(list[i], list2[i])
 doc.save(dosyayolu + '5.doc')
 return 1


def replace_string6(filename='8.ARABULUCULUK BÜRO ÖDEME ÜST YAZI.docx'):
    doc = Document(filename)
    for p in doc.paragraphs:
        inline = p.runs
        for j in range(0, len(inline)):
            for i in range(0, len(list)):
               p.text = p.text.replace(list[i], list2[i])
    doc.save(dosyayolu + '6.doc')
    return 1


def replace_string7(filename='7.ARABULUCULUK SON TUTANAK GÖNDERİMİ ÜST YAZI.docx'):
 doc = Document(filename)
 for p in doc.paragraphs:
    inline = p.runs
    for j in range(0, len(inline)):
        for i in range(0, len(list)):
           p.text = p.text.replace(list[i], list2[i])
 doc.save(dosyayolu + '7.doc')
 return 1


def replace_string8(filename='6.SON TUTANAK.docx'):
    doc = Document(filename)
    for p in doc.paragraphs:
        inline = p.runs
        for j in range(0, len(inline)):
            for i in range(0, len(list)):
               p.text = p.text.replace(list[i], list2[i])
    doc.save(dosyayolu + '8.doc')
    return 1


def replace_string9(filename='9.ÖRNEK MAKBUZ.docx'):
    doc = Document(filename)
    for p in doc.paragraphs:
        inline = p.runs
        for j in range(0, len(inline)):
            for i in range(0, len(list)):
               p.text = p.text.replace(list[i], list2[i])
    doc.save(dosyayolu + '9.doc')
    return 1


def replace_string10(filename='10.ÜCRET SÖZLEŞMESİ.docx'):
    doc = Document(filename)
    for p in doc.paragraphs:
        inline = p.runs
        for j in range(0, len(inline)):
            for i in range(0, len(list)):
               p.text = p.text.replace(list[i], list2[i])
    doc.save(dosyayolu + '10.doc')
    return 1


def replace_string11(filename='12.DAVET MEKTUBU.docx'):
    doc = Document(filename)
    for p in doc.paragraphs:
        inline = p.runs
        for j in range(0, len(inline)):
            for i in range(0, len(list)):
               p.text = p.text.replace(list[i], list2[i])
    doc.save(dosyayolu + '11.doc')
    return 1


replace_string1()
replace_string2()
replace_string3()
replace_string4()
replace_string5()
replace_string6()
replace_string7()
replace_string8()
replace_string9()
replace_string10()
replace_string11()

import aspose.words as aw
doc1 = aw.Document(dosyayolu + '1.doc')
doc2 = aw.Document(dosyayolu + '2.doc')
doc3 = aw.Document(dosyayolu + '3.doc')
doc4 = aw.Document(dosyayolu + '4.doc')
doc5 = aw.Document(dosyayolu + '5.doc')
doc6 = aw.Document(dosyayolu + '6.doc')
doc7 = aw.Document(dosyayolu + '7.doc')
doc8 = aw.Document(dosyayolu + '8.doc')
doc9 = aw.Document(dosyayolu + '9.doc')
doc10 = aw.Document(dosyayolu + '10.doc')
doc11 = aw.Document(dosyayolu + '11.doc')
doc1.save(dosyayolu+"1. Davet Mektubu.rtf")
doc2.save(dosyayolu+"2. Bilgilendirme Tutanağı.rtf")
doc3.save(dosyayolu+"3. İlk Oturum Tutanağı.rtf")
doc4.save(dosyayolu+"4. İkinci Oturum.rtf")
doc5.save(dosyayolu+"5. Anlaşma Belgesi.rtf")
doc6.save(dosyayolu+"6. Son Tutanak.rtf")
doc7.save(dosyayolu+"7. Arabulucukuk Son Tutanak.rtf")
doc8.save(dosyayolu+"8. Büro Ödeme ÜSt.rtf")
doc9.save(dosyayolu+"9. Örnek Makbuz.rtf")
doc10.save(dosyayolu+"10. Ücret Sözleşmesi.rtf")
doc11.save(dosyayolu+"12. Davet Mektubu.rtf")

